import streamlit as st
import os
import pickle
from datetime import datetime
from typing import List
import logging
import re
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import PyPDF2
import docx
from io import BytesIO
import time

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Page config
st.set_page_config(
    page_title="NoteChamp - AI Study Assistant",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Session states
if 'app_launched' not in st.session_state:
    st.session_state.app_launched = False
if 'welcome_shown' not in st.session_state:
    st.session_state.welcome_shown = False
if 'chatbot' not in st.session_state:
    st.session_state.chatbot = None
if 'current_subject' not in st.session_state:
    st.session_state.current_subject = None
if 'last_response' not in st.session_state:
    st.session_state.last_response = None
if 'last_question' not in st.session_state:
    st.session_state.last_question = None

# CSS
def load_css():
    st.markdown("""
    <style>
    /* Hide default Streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Animations */
    @keyframes fadeIn { from {opacity:0;} to {opacity:1;} }
    @keyframes zoomBounce { 0%{transform:scale(0);opacity:0;} 60%{transform:scale(1.1);} 80%{transform:scale(0.95);} 100%{transform:scale(1);opacity:1;} }
    @keyframes slideUp { from {opacity:0;transform:translateY(30px);} to {opacity:1;transform:translateY(0);} }
    @keyframes slideInRight { from {opacity:0;transform:translateX(40px);} to {opacity:1;transform:translateX(0);} }
    @keyframes slideInLeft { from {opacity:0;transform:translateX(-40px);} to {opacity:1;transform:translateX(0);} }
    @keyframes fadeOut { to {opacity:0;visibility:hidden;} }
    
    /* Fullscreen Welcome/Launch screens */
    .fullscreen, .fullscreen-welcome, .launch-overlay {
        position: fixed; top:0; left:0; width:100vw; height:100vh;
        background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
        display:flex; justify-content:center; align-items:center;
        z-index:10000; animation:fadeIn 0.5s ease-in;
        color:#f0f0f0;
    }
    .logo-text, .welcome-text { font-size:72px; font-weight:800; text-shadow: 0 6px 30px rgba(0,0,0,0.4); animation:zoomBounce 1.2s ease-out; color:#f0f0f0; }
    .launch-logo { font-size:120px; animation: zoomBounce 1s ease-out 0.3s both; margin-bottom:20px; }
    .launch-title { font-size:64px; font-weight:bold; color:#67d1ff; letter-spacing:8px; animation:slideUp 0.8s ease-out 1s both; text-shadow:0 4px 20px rgba(0,0,0,0.3); }
    .launch-subtitle { font-size:20px; color:#cbd5e1; animation:slideUp 0.8s ease-out 1.3s both; margin-top:10px; }

    /* App header */
    .app-header { text-align:center; padding:20px 0 30px 0; background:linear-gradient(135deg,#1e1e1e,#2d2d2d); color:#f0f0f0; border-radius:0 0 30px 30px; margin:-60px -100px 30px -100px; box-shadow:0 4px 20px rgba(0,0,0,0.5); }
    .app-header h1 { font-size:48px; margin:0; font-weight:700; letter-spacing:2px; }
    .app-header p { font-size:18px; margin:10px 0 0 0; opacity:0.8; }

    /* Sidebar dark theme */
    section[data-testid="stSidebar"] {
        background-color: #1e1e1e !important;
        color: #f0f0f0 !important;
    }
    section[data-testid="stSidebar"] .stTextInput>div>div>input,
    section[data-testid="stSidebar"] .stButton>button,
    section[data-testid="stSidebar"] select {
        background-color: #2d2d2d !important;
        color: #f0f0f0 !important;
        border: 1px solid #444 !important;
    }
    section[data-testid="stSidebar"] .stButton>button:hover {
        background-color: #444 !important;
    }

    /* Chat bubbles */
    .user-message { background:#3b82f6; color:white; padding:16px 22px; border-radius:22px 22px 5px 22px; margin:15px 0 15px auto; max-width:75%; box-shadow:0 4px 12px rgba(0,0,0,0.3); animation:slideInRight 0.4s ease-out; font-size:15px; line-height:1.6; }
    .ai-message { background:#2d2d2d; color:#f0f0f0; padding:24px; border-radius:22px 22px 22px 5px; margin:15px auto 15px 0; max-width:80%; box-shadow:0 4px 16px rgba(0,0,0,0.3); border-left:4px solid #67d1ff; animation:slideInLeft 0.4s ease-out; }

    .response-header { display:flex; align-items:center; margin-bottom:16px; padding-bottom:12px; border-bottom:2px solid #444; }
    .response-icon { font-size:28px; margin-right:12px; }
    .response-type { font-size:20px; font-weight:600; color:#67d1ff; }
    .question-badge { background:#2d2d2d; color:#67d1ff; padding:10px 16px; border-radius:12px; font-weight:600; margin-bottom:16px; display:inline-block; border:2px solid #444; }
    .answer-content { line-height:1.9; font-size:15px; color:#f0f0f0; }
    .success-badge { display:inline-block; padding:6px 14px; background:#16a34a; color:white; border-radius:16px; font-size:13px; font-weight:600; margin-top:12px; box-shadow:0 2px 8px rgba(22,163,52,0.3); }

    /* Feature cards */
    .feature-card { background:#2d2d2d; color:white; padding:30px; border-radius:20px; box-shadow:0 4px 16px rgba(0,0,0,0.3); margin:20px; display:inline-block; text-align:left; transition:all 0.3s ease; }
    .feature-card:hover { transform:translateY(-5px); box-shadow:0 8px 24px rgba(102,126,234,0.2); }
    .feature-icon { font-size:48px; margin-bottom:15px; }
    .feature-title { font-size:24px; font-weight:600; color:#67d1ff; margin-bottom:10px; }
    .feature-desc { font-size:15px; color:#cbd5e1; line-height:1.6; }

    /* Inputs/buttons */
    .stTextInput>div>div>input { border-radius:12px !important; border:2px solid #444 !important; padding:12px 16px !important; font-size:15px !important; background:#2d2d2d !important; color:#f0f0f0 !important; }
    .stTextInput>div>div>input:focus { border-color:#67d1ff !important; box-shadow:0 0 0 3px rgba(102,209,255,0.2) !important; }
    .stButton>button { border-radius:12px !important; padding:10px 20px !important; font-weight:600 !important; transition:all 0.3s ease !important; border:none !important; box-shadow:0 2px 8px rgba(0,0,0,0.3) !important; background:#3b3b3b; color:#f0f0f0; }
    .stButton>button:hover { transform:translateY(-2px) !important; box-shadow:0 4px 12px rgba(0,0,0,0.5) !important; background:#444 !important; }

    /* Download badges */
    .doc-badge { background:#3b82f6; color:white; padding:8px 16px; border-radius:20px; font-size:13px; display:inline-block; margin:5px 5px 5px 0; font-weight:600; box-shadow:0 2px 8px rgba(59,130,246,0.3); }
    </style>
    """, unsafe_allow_html=True)


def show_logo():
    st.markdown('<div class="app-logo"><span class="app-logo-icon">📚</span><span>NoteChamp</span></div>', unsafe_allow_html=True)

# Document processor
class SimpleDocProcessor:
    @staticmethod
    def extract_text(file) -> str:
        try:
            ext = file.name.lower().split('.')[-1]
            if ext == 'pdf':
                pdf = PyPDF2.PdfReader(file)
                return "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
            elif ext == 'docx':
                doc = docx.Document(file)
                return "\n".join([p.text for p in doc.paragraphs])
            elif ext == 'txt':
                content = file.read()
                return content.decode('utf-8') if isinstance(content, bytes) else content
            return ""
        except Exception as e:
            logger.error(f"Error: {e}")
            return ""

    @staticmethod
    def chunk_text(text: str, size: int = 300) -> List[str]:
        words = text.split()
        return [" ".join(words[i:i+size]) for i in range(0, len(words), size) if words[i:i+size]]

# Vector store
class SimpleVectorStore:
    def __init__(self):
        self.model = None
        self.index = None
        self.chunks = []

    def load_model(self):
        if self.model is None:
            self.model = SentenceTransformer("all-MiniLM-L6-v2")

    def add_documents(self, chunks: List[str]):
        if not chunks: return
        self.load_model()
        self.chunks.extend(chunks)
        emb = self.model.encode(self.chunks)
        self.index = faiss.IndexFlatIP(emb.shape[1])
        faiss.normalize_L2(emb)
        self.index.add(emb)

    def search(self, query: str, k: int = 5) -> List[str]:
        if not self.chunks or self.index is None: return []
        self.load_model()
        qe = self.model.encode([query])
        faiss.normalize_L2(qe)
        scores, indices = self.index.search(qe, min(k, len(self.chunks)))
        return [self.chunks[i] for s,i in zip(scores[0], indices[0]) if i < len(self.chunks) and s>0.05]

# Response formatter
class ResponseFormatter:
    @staticmethod
    def clean(text: str) -> str:
        text = re.sub(r'MUQuestionPapers?\.com\s*Page\s*\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Q\.?P\.?\s*Code\s*[–-]?\s*\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\bQ\.?\s*\d+\s*[a-z]?\)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b\d+M\b', '', text)
        text = re.sub(r'\(\s*\d+\s*marks?\s*\)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\b(Answer|Solution|Ans)\s*:\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    @staticmethod
    def extract_sentences(text: str, max_s: int = 6) -> List[str]:
        sents = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sents if len(s.strip())>30 and s[0].isupper() and s[-1] in '.!?'][:max_s]

    @staticmethod
    def format_response(question: str, raw: str, qtype: str) -> dict:
        cleaned = ResponseFormatter.clean(raw)
        sentences = ResponseFormatter.extract_sentences(cleaned, 6)
        if not sentences: return None
        
        if qtype in ["List/Types","Process/Steps","Advantages/Disadvantages"]:
            content = '<br><br>'.join([f"• {s.rstrip('.')}" for s in sentences])
        else:
            paras = []
            curr = []
            for i, s in enumerate(sentences):
                curr.append(s)
                if len(curr)==2 or i==len(sentences)-1:
                    paras.append(' '.join(curr))
                    curr=[]
            content = '<br><br>'.join(paras)
        
        return {'question_type': qtype, 'content': content, 'sentences': sentences, 'raw': raw}

# AI helper
class AIHelper:
    @staticmethod
    def get_qtype(q: str) -> str:
        ql = q.lower()
        if any(w in ql for w in ['what is','define','definition']): return "Definition"
        if any(w in ql for w in ['explain','describe','discuss']): return "Explanation"
        if any(w in ql for w in ['how','steps','process','method']): return "Process/Steps"
        if any(w in ql for w in ['why','reason','cause']): return "Reasoning"
        if any(w in ql for w in ['compare','difference','vs']): return "Comparison"
        if any(w in ql for w in ['list','types','kinds']): return "List/Types"
        if any(w in ql for w in ['advantage','disadvantage','pros','cons']): return "Advantages/Disadvantages"
        return "Answer"

    @staticmethod
    def get_icon(qtype: str) -> str:
        icons = {"Definition":"📖","Explanation":"💡","Process/Steps":"🔄","Reasoning":"🤔","Comparison":"⚖️","List/Types":"📋","Advantages/Disadvantages":"✓✗","Answer":"💬"}
        return icons.get(qtype, "💬")

    @staticmethod
    def generate(q: str, ctx: str) -> dict:
        if not ctx.strip(): return None
        qtype = AIHelper.get_qtype(q)
        result = ResponseFormatter.format_response(q, ctx, qtype)
        if result: result['icon'] = AIHelper.get_icon(qtype)
        return result

# Export
class ExportHelper:
    @staticmethod
    def to_txt(q: str, ans: str) -> BytesIO:
        clean = re.sub('<br>', '\n', ans)
        clean = re.sub('<[^<]+?>', '', clean)
        buf = BytesIO()
        buf.write(f"NoteChamp Answer\n\nQ: {q}\n\nA:\n{clean}".encode('utf-8'))
        buf.seek(0)
        return buf

    @staticmethod
    def to_docx(q: str, ans: str) -> BytesIO:
        from docx import Document
        from docx.shared import Pt, RGBColor
        clean = re.sub('<br>', '\n', ans)
        clean = re.sub('<[^<]+?>', '', clean)
        doc = Document()
        t = doc.add_heading('NoteChamp Answer', 0)
        t.alignment = 1
        qp = doc.add_paragraph()
        qr = qp.add_run(f'Q: {q}')
        qr.bold = True
        qr.font.size = Pt(14)
        qr.font.color.rgb = RGBColor(102,126,234)
        doc.add_paragraph()
        ap = doc.add_paragraph()
        ar = ap.add_run('Answer:')
        ar.bold = True
        ar.font.size = Pt(12)
        doc.add_paragraph(clean)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

# Chatbot
class StudyChatbot:
    def __init__(self):
        self.subjects = {}
        self.data_dir = "chatbot_data"
        os.makedirs(self.data_dir, exist_ok=True)

    def create_subject(self, name: str):
        if name not in self.subjects:
            self.subjects[name] = {
                'vector_store': SimpleVectorStore(),
                'chat_history': [],
                'documents': [],
                'created_at': datetime.now()
            }
            self.save()

    def add_documents(self, subj: str, files):
        if subj not in self.subjects:
            self.create_subject(subj)
        chunks = []
        existing = [d['name'] for d in self.subjects[subj]['documents']]
        new = 0
        for f in files:
            if f.name in existing: continue
            txt = SimpleDocProcessor.extract_text(f)
            if txt:
                ch = SimpleDocProcessor.chunk_text(txt)
                chunks.extend(ch)
                self.subjects[subj]['documents'].append({'name':f.name,'uploaded_at':datetime.now(),'chunks':len(ch)})
                new += 1
        if chunks:
            self.subjects[subj]['vector_store'].add_documents(chunks)
            self.save()
        return len(chunks), new

    def chat(self, subj: str, q: str) -> dict:
        if subj not in self.subjects: return None
        vs = self.subjects[subj]['vector_store']
        if not vs.chunks: return None
        chunks = vs.search(q, 5)
        cleaned = [ResponseFormatter.clean(c) for c in chunks if len(c.split())>15]
        ctx = " ".join(cleaned)
        result = AIHelper.generate(q, ctx)
        if result:
            self.subjects[subj]['chat_history'].append({'question':q,'response':result,'timestamp':datetime.now()})
            self.save()
        return result

    def save(self):
        try:
            meta = {}
            for n, d in self.subjects.items():
                meta[n] = {'chat_history':d['chat_history'],'documents':d['documents'],'created_at':d['created_at'],'chunks':d['vector_store'].chunks}
            with open(os.path.join(self.data_dir, 'subjects.pkl'), 'wb') as f:
                pickle.dump(meta, f)
        except Exception as e:
            logger.error(f"Save error: {e}")

    def load(self):
        try:
            path = os.path.join(self.data_dir, 'subjects.pkl')
            if os.path.exists(path):
                with open(path, 'rb') as f:
                    meta = pickle.load(f)
                for n, d in meta.items():
                    self.subjects[n] = {'vector_store':SimpleVectorStore(),'chat_history':d['chat_history'],'documents':d['documents'],'created_at':d['created_at']}
                    if 'chunks' in d and d['chunks']:
                        self.subjects[n]['vector_store'].add_documents(d['chunks'])
        except Exception as e:
            logger.error(f"Load error: {e}")

# Display
def display_response(q: str, resp: dict):
    st.markdown(f"""
    <div class="ai-message">
        <div class="response-header">
            <span class="response-icon">{resp['icon']}</span>
            <span class="response-type">{resp['question_type']}</span>
        </div>
        <div class="question-badge">Q: {q}</div>
        <div class="answer-content">{resp['content']}</div>
        <div class="success-badge">✓ From study materials</div>
    </div>
    """, unsafe_allow_html=True)

# Main
def main():
   # Animation sequence
    if not st.session_state.welcome_shown:
        placeholder = st.empty()
        with placeholder.container():
            st.markdown("""
            <div class="fullscreen" id="welcomeScreen">
                <div class="logo-text">Welcome to NoteChamp</div>
            </div>
            <script>
                setTimeout(function() {
                    var screen = document.getElementById('welcomeScreen');
                    if (screen) {
                        screen.style.animation = 'fadeOut 0.8s ease-out forwards';
                        setTimeout(function(){ screen.remove(); }, 800);
                    }
                }, 1500);
            </script>
            """, unsafe_allow_html=True)
        st.session_state.welcome_shown = True
    # no st.rerun() here! just let the rest of the app load

    elif not st.session_state.app_launched:
        placeholder = st.empty()
        with placeholder.container():
            st.markdown("""
            <div class="launch-overlay" id="launch">
                <div class="launch-logo">📚</div>
                <div class="launch-title">NOTECHAMP</div>
                <div class="launch-subtitle">Your AI Study Companion</div>
            </div>
            <script>
                setTimeout(function() {
                    var l = document.getElementById('launch');
                    if (l) {
                        l.style.animation = 'fadeOut 0.5s ease-out forwards';
                        setTimeout(function(){ l.remove(); }, 500);
                    }
                }, 1000);
            </script>
            """, unsafe_allow_html=True)
        st.session_state.app_launched = True
# again, no st.rerun()


    
    # Now render main UI
    load_css()
    show_logo()
    
    if st.session_state.chatbot is None:
        st.session_state.chatbot = StudyChatbot()
        st.session_state.chatbot.load()
    
    cb = st.session_state.chatbot
    
    st.markdown('<div class="app-header"><h1>📚 NoteChamp</h1><p>Your AI Study Companion</p></div>', unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.markdown("### 📚 Subjects")
        new = st.text_input("Create subject", placeholder="e.g., OS", label_visibility="collapsed")
        if st.button("Create", use_container_width=True, type="primary"):
            if new.strip():
                cb.create_subject(new.strip())
                st.session_state.current_subject = new.strip()
                st.success("Created!")
                time.sleep(0.5)
                st.rerun()
        
        if cb.subjects:
            st.markdown("---")
            subs = list(cb.subjects.keys())
            idx = 0
            if st.session_state.current_subject in subs:
                idx = subs.index(st.session_state.current_subject)
            
            sel = st.selectbox("", subs, index=idx, label_visibility="collapsed")
            if sel != st.session_state.current_subject:
                st.session_state.current_subject = sel
                st.rerun()
            
            if sel:
                data = cb.subjects[sel]
                c1, c2 = st.columns(2)
                c1.metric("Docs", len(data['documents']))
                c2.metric("Chats", len(data['chat_history']))
        
        if st.session_state.current_subject:
            st.markdown("---")
            st.markdown("### 📤 Upload")
            files = st.file_uploader("", accept_multiple_files=True, type=['pdf','txt','docx'], label_visibility="collapsed")
            if files and st.button("Upload", use_container_width=True, type="primary"):
                with st.spinner("Processing..."):
                    ch, cnt = cb.add_documents(st.session_state.current_subject, files)
                    if cnt > 0:
                        st.success(f"Added {cnt} files!")
                        time.sleep(0.8)
                        st.rerun()
    
    # Main
    if st.session_state.current_subject:
        st.markdown(f"## 💬 {st.session_state.current_subject}")
        
        q = st.text_input("Ask", placeholder="Your question...", label_visibility="collapsed")
        
        if st.button("Ask", type="primary") and q.strip():
            st.markdown(f'<div class="user-message">{q.strip()}</div>', unsafe_allow_html=True)
            
            with st.spinner(""):
                result = cb.chat(st.session_state.current_subject, q.strip())
                if result:
                    st.session_state.last_response = result
                    st.session_state.last_question = q.strip()
                    st.rerun()
        
        if st.session_state.last_response:
            display_response(st.session_state.last_question, st.session_state.last_response)
            
            st.markdown("### 📥 Export")
            c1, c2, c3 = st.columns(3)
            
            with c1:
                txt = ExportHelper.to_txt(st.session_state.last_question, st.session_state.last_response['content'])
                st.download_button("📄 TXT", txt, f"answer_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt", "text/plain", use_container_width=True)
            
            with c2:
                docx = ExportHelper.to_docx(st.session_state.last_question, st.session_state.last_response['content'])
                st.download_button("📄 DOCX", docx, f"answer_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    
    # ===== Main UI Section =====
    else:
        st.markdown("""
            
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
